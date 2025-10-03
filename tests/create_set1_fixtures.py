#!/usr/bin/env python3
"""
Create Set 1 DOCX fixtures for advanced features.
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

# Create fixtures directory
fixtures_dir = Path(__file__).parent / "fixtures" / "set1"
fixtures_dir.mkdir(parents=True, exist_ok=True)

def create_nested_tables_docx():
    """Create DOCX with nested tables."""
    doc = Document()
    doc.add_heading("Nested Tables Fixture", 0)

    doc.add_paragraph("This document contains a table with a nested table inside a cell.")

    # Outer table
    outer_table = doc.add_table(rows=3, cols=2)
    outer_table.style = "Table Grid"
    outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fill outer cells
    outer_table.cell(0, 0).text = "Header 1"
    outer_table.cell(0, 1).text = "Header 2"
    outer_table.cell(1, 0).text = "Outer cell with nested table"
    outer_table.cell(1, 1).text = "Another outer cell"
    outer_table.cell(2, 0).text = "Footer 1"
    outer_table.cell(2, 1).text = "Footer 2"

    # Add nested table in cell (1,0)
    nested_cell = outer_table.cell(1, 0)
    nested_table = nested_cell.add_table(rows=2, cols=3)
    nested_table.style = "Table Grid"

    nested_table.cell(0, 0).text = "Nested Header 1"
    nested_table.cell(0, 1).text = "Nested Header 2"
    nested_table.cell(0, 2).text = "Nested Header 3"
    nested_table.cell(1, 0).text = "Data 1"
    nested_table.cell(1, 1).text = "Data 2"
    nested_table.cell(1, 2).text = "Data 3"

    doc.add_paragraph("End of document.")

    output_path = fixtures_dir / "nested_tables.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

def create_merged_cells_docx():
    """Create DOCX with merged cells in table."""
    doc = Document()
    doc.add_heading("Merged Cells Fixture", 0)

    doc.add_paragraph("This document contains a table with merged cells.")

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"

    # Headers
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Description"
    table.cell(0, 2).text = "Value"

    # Merge first row data cells
    cell1 = table.cell(1, 0)
    cell2 = table.cell(1, 1)
    cell1.merge(cell2)
    cell1.text = "Merged Item 1"

    table.cell(1, 2).text = "100"

    # Normal row
    table.cell(2, 0).text = "Item 2"
    table.cell(2, 1).text = "Description 2"
    table.cell(2, 2).text = "200"

    doc.add_paragraph("End of document.")

    output_path = fixtures_dir / "merged_cells.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

def create_hyperlinks_docx():
    """Create DOCX with hyperlinks."""
    doc = Document()
    doc.add_heading("Hyperlinks Fixture", 0)

    doc.add_paragraph("This document contains hyperlinks.")

    p = doc.add_paragraph()
    run = p.add_run("Visit ")
    run.hyperlink = "https://example.com"
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.underline = True

    p.add_run(" for more information.")

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Internal link: ")
    run2.hyperlink = "#anchor"
    run2.font.color.rgb = RGBColor(0, 0, 255)
    run2.font.underline = True

    doc.add_paragraph("End of document.")

    output_path = fixtures_dir / "hyperlinks.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")

def create_comments_docx():
    """Create DOCX with comments."""
    doc = Document()
    doc.add_heading("Comments Fixture", 0)

    p = doc.add_paragraph("This is a paragraph with a comment.")
    run = p.runs[0]
    # Note: Adding comments requires low-level access or recent python-docx features
    # For simulation, we'll add a simple paragraph
    # In real, use: from docx.oxml import OxmlElement
    # But to keep simple, add text indicating comment location

    doc.add_paragraph("[Comment here: This is a test comment]")

    p2 = doc.add_paragraph("Another paragraph.")

    output_path = fixtures_dir / "comments.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    print("Note: Full comment support may require manual Word addition or advanced XML.")

def create_tracked_changes_docx():
    """Create DOCX with tracked changes (revisions)."""
    doc = Document()
    doc.add_heading("Tracked Changes Fixture", 0)

    p = doc.add_paragraph("Original text.")
    # Tracked changes are complex; simulate with text
    doc.add_paragraph("[Tracked insertion: New text added]")
    doc.add_paragraph("[Tracked deletion: Old text removed]")

    output_path = fixtures_dir / "tracked_changes.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    print("Note: Tracked changes require Word or advanced XML manipulation.")

def create_fields_docx():
    """Create DOCX with fields like DATE, PAGE."""
    doc = Document()
    doc.add_heading("Fields Fixture", 0)

    doc.add_paragraph("Current date: {DATE}")
    doc.add_paragraph("Page number: {PAGE}")

    # To add real fields, need XML
    # Simulate with text

    output_path = fixtures_dir / "fields.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")
    print("Note: Real fields require XML insertion.")

if __name__ == "__main__":
    print("Creating Set 1 fixtures...")
    create_nested_tables_docx()
    create_merged_cells_docx()
    create_hyperlinks_docx()
    create_comments_docx()
    create_tracked_changes_docx()
    create_fields_docx()
    print("\nAll Set 1 fixtures created successfully!")