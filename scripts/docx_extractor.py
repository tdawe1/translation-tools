from docx import Document
from typing import List, Tuple, Dict, Any
from .base_extractor import BaseExtractor

class DocxExtractor(BaseExtractor):
    def extract(self, file_path: str, page_range: Optional[str] = None) -> List[Tuple[str, int, str, Dict[str, Any]]]:
        doc = Document(file_path)
        blocks = []
        
        # Extract paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                layout = {
                    'type': 'paragraph',
                    'style': para.style.name if para.style else 'Normal',
                    'font_size': para.runs[0].font.size.pt if para.runs else 12.0,
                    'bold': any(run.bold for run in para.runs) if para.runs else False,
                    'italic': any(run.italic for run in para.runs) if para.runs else False
                }
                blocks.append(('body', i, text, layout))
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        text = para.text.strip()
                        if text:
                            section = f'table_{table_idx}_row_{row_idx}_cell_{cell_idx}'
                            layout = {
                                'type': 'table_cell',
                                'table_idx': table_idx,
                                'row_idx': row_idx,
                                'cell_idx': cell_idx,
                                'font_size': para.runs[0].font.size.pt if para.runs else 11.0
                            }
                            blocks.append((section, para_idx, text, layout))
        
        return blocks