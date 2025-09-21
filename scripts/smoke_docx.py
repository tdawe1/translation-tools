#!/usr/bin/env python3
"""
Smoke test for translate_docx.py CLI
Runs on cli_sample.docx, checks outputs.
For local: set DRY_RUN=1 to use --dry-run (skips translation, expects JP >0)
For CI: real translation, expects JP ==0
"""
import csv
import json
import os
import subprocess
import sys
from pathlib import Path

from docx import Document


def main():
    dry_run = os.getenv('DRY_RUN', '0') == '1'
    fixture_dir = Path(__file__).parent.parent / "tests"
    input_file = fixture_dir / "cli_sample.docx"
    if not input_file.exists():
        print(f"Error: Fixture {input_file} not found", file=sys.stderr)
        return 1
    tmp_dir = Path.cwd() / "tmp_smoke_docx"
    tmp_dir.mkdir(exist_ok=True)
    output_file = tmp_dir / "output.docx"
    csv_file = tmp_dir / "bilingual.csv"
    audit_file = tmp_dir / "audit.json"
    cmd = [
        sys.executable, "scripts/translate_docx.py",
        "--in", str(input_file),
        "--out", str(output_file),
        "--bilingual-csv",
        "--json-audit",
        "--model", "gpt-4o-mini",
        "--batch", "5",
        "--no-cache",
        "--no-backup"
    ]
    if dry_run:
        cmd += ["--dry-run"]
    result = subprocess.run(cmd, cwd=Path(__file__).parent, capture_output=True, text=True)
    if result.returncode != 0:
        print("Translation CLI failed:", file=sys.stderr)
        print(result.stderr, file=sys.stderr)
        return 1
    if not output_file.exists():
        print("Output file not created", file=sys.stderr)
        return 1
    # Check CSV
    if not csv_file.exists():
        print("CSV not created", file=sys.stderr)
        return 1
    try:
        with open(csv_file, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            rows = list(reader)
            if len(rows) < 2:
                print("CSV has no data rows", file=sys.stderr)
                return 1
            if rows[0] != ['Segment ID', 'Original Japanese', 'Translated English', 'Context']:
                print("CSV header mismatch", file=sys.stderr)
                return 1
    except Exception as e:
        print(f"CSV validation failed: {e}", file=sys.stderr)
        return 1
    # Check audit JSON
    if not audit_file.exists():
        print("Audit JSON not created", file=sys.stderr)
        return 1
    try:
        with open(audit_file, 'r', encoding='utf-8') as f:
            audit = json.load(f)
            if "translation_info" not in audit or "translation_stats" not in audit:
                print("Audit missing required keys", file=sys.stderr)
                return 1
            total_segments = audit["translation_stats"]["total_segments"]
            if total_segments == 0:
                print("No segments found", file=sys.stderr)
                return 1
    except Exception as e:
        print(f"Audit validation failed: {e}", file=sys.stderr)
        return 1
    # Check XML parity using python-docx
    try:
        doc_in = Document(input_file)
        doc_out = Document(output_file)
        para_in = len(doc_in.paragraphs)
        para_out = len(doc_out.paragraphs)
        if abs(para_in - para_out) > 1:  # allow minor diff
            print(f"Paragraph count mismatch: {para_in} vs {para_out}", file=sys.stderr)
            return 1
    except Exception as e:
        print(f"Structure check failed: {e}", file=sys.stderr)
        return 1
    # Run JP audit
    audit_cmd = [sys.executable, "scripts/audit_docx_jp_count.py", str(output_file)]
    audit_result = subprocess.run(audit_cmd, cwd=Path(__file__).parent, capture_output=True, text=True)
    if dry_run:
        if audit_result.returncode == 0:
            print("JP audit passed unexpectedly in dry-run mode (should have JP chars)", file=sys.stderr)
            return 1
    else:
        if audit_result.returncode != 0:
            print("JP audit failed in full mode", file=sys.stderr)
            print(audit_result.stdout, file=sys.stderr)
            return 1
    print("All smoke checks passed")
    return 0

if __name__ == "__main__":
    sys.exit(main())
