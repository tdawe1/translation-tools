#!/usr/bin/env python3
"""
DOCX translation helper for manual/local workflows.

Usage:
  # Step 1: extract unique Japanese segments to a template JSON
  python scripts/manual_docx_translation.py prepare \
      --input inputs/source.docx \
      --template translations/source_template.json

  # Step 2: fill in the "en" fields (manually or via Codex assistance)

  # Step 3: apply translations to create an English DOCX
  python scripts/manual_docx_translation.py apply \
      --input inputs/source.docx \
      --translations translations/source_translations.json \
      --output outputs/source_en.docx
"""
from __future__ import annotations

import argparse
import io
import json
import zipfile
from pathlib import Path
from typing import Dict, Iterable, List, Tuple
import xml.etree.ElementTree as ET

# Ensure we import the external 'python-docx' library, not local 'scripts/docx'
import sys, os
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR in sys.path:
    try:
        sys.path.remove(_SCRIPT_DIR)
    except ValueError:
        pass
from docx import Document


W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def collect_segments(docx_path: Path) -> List[str]:
    """Collect unique non-empty text segments (paragraphs + table cells)."""
    document = Document(docx_path)
    segments: List[str] = []
    seen: set[str] = set()

    def add_text(text: str) -> None:
        if text is None:
            return
        if not text.strip():
            return
        if text not in seen:
            segments.append(text)
            seen.add(text)

    for paragraph in document.paragraphs:
        add_text(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    add_text(paragraph.text)

    return segments


def paragraph_text(paragraph: ET.Element) -> str:
    return "".join((t.text or "") for t in paragraph.iter(f"{W_NS}t"))


def replace_paragraph_text(paragraph: ET.Element, replacement: str) -> None:
    text_nodes = [t for t in paragraph.iter(f"{W_NS}t")]
    if not text_nodes:
        return
    text_nodes[0].text = replacement
    for node in text_nodes[1:]:
        node.text = ""


def translate_docx(
    input_path: Path,
    output_path: Path,
    translation_map: Dict[str, str],
) -> List[str]:
    """Apply translations stored in translation_map. Returns list of missing JP strings."""
    missing: List[str] = []

    with zipfile.ZipFile(input_path, "r") as src_zip, zipfile.ZipFile(output_path, "w") as dst_zip:
        for info in src_zip.infolist():
            data = src_zip.read(info.filename)
            if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                root = ET.fromstring(data)
                for paragraph in root.iter(f"{W_NS}p"):
                    src_text = paragraph_text(paragraph)
                    lookup = src_text.strip()
                    if not lookup:
                        continue
                    if src_text in translation_map:
                        replace_paragraph_text(paragraph, translation_map[src_text])
                    elif lookup in translation_map:
                        replace_paragraph_text(paragraph, translation_map[lookup])
                    else:
                        missing.append(src_text)
                buffer = io.BytesIO()
                ET.ElementTree(root).write(buffer, encoding="utf-8", xml_declaration=True)
                data = buffer.getvalue()
            dst_zip.writestr(info, data)

    return missing


def load_translation_map(translations_path: Path) -> Dict[str, str]:
    data = json.loads(translations_path.read_text(encoding="utf-8"))
    mapping: Dict[str, str] = {}
    for entry in data:
        jp = entry.get("jp", "")
        en = entry.get("en", "")
        if not jp.strip():
            continue
        if not en.strip():
            raise ValueError(f"Missing English translation for segment: {jp}")
        mapping[jp] = en
    return mapping


def write_template(segments: Iterable[str], template_path: Path) -> None:
    template_path.parent.mkdir(parents=True, exist_ok=True)
    template = [{"jp": seg, "en": ""} for seg in segments]
    template_path.write_text(json.dumps(template, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Manual DOCX translation helper")
    subparsers = parser.add_subparsers(dest="command", required=True)

    prep = subparsers.add_parser("prepare", help="Extract unique segments to a translation template")
    prep.add_argument("--input", required=True, type=Path, help="Source DOCX file")
    prep.add_argument("--template", required=True, type=Path, help="Output path for template JSON")

    apply = subparsers.add_parser("apply", help="Apply translations from JSON to DOCX")
    apply.add_argument("--input", required=True, type=Path, help="Source DOCX file")
    apply.add_argument("--translations", required=True, type=Path, help="JSON file with jp/en pairs")
    apply.add_argument("--output", required=True, type=Path, help="Output DOCX path")

    return parser.parse_args()


def main() -> None:
    args = parse_args()

    if args.command == "prepare":
        segments = collect_segments(args.input)
        write_template(segments, args.template)
        print(f"Wrote {len(segments)} unique segments to {args.template}")
        return

    if args.command == "apply":
        args.output.parent.mkdir(parents=True, exist_ok=True)
        translation_map = load_translation_map(args.translations)
        missing = translate_docx(args.input, args.output, translation_map)
        if missing:
            raise SystemExit(
                "Translation completed with missing segments:\n"
                + "\n".join(missing[:20])
                + ("\n..." if len(missing) > 20 else "")
            )
        print(f"Translated DOCX written to {args.output}")
        return


if __name__ == "__main__":
    main()
